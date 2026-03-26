"""
KSB Coursework Marker â€” LangGraph Version

Streamlit UI wired to the compiled LangGraph StateGraph.
Flow: Upload DOCX/PDF -> Ingest -> graph.stream() -> Display results
"""
import streamlit as st
import tempfile
import time
import logging
import shutil
import json
import csv
import base64
import hashlib
import html
from io import StringIO
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.document_processing import DocxProcessor, PDFProcessor, ImageProcessor
from src.chunking import SmartChunker
from src.embeddings import Embedder
from src.vector_store import ChromaStore
from src.llm import OllamaClient
from src.criteria import get_module_criteria, get_available_modules
from src.brief import get_default_brief, parse_uploaded_brief

from config import (
    OLLAMA_BASE_URL, OLLAMA_MODEL, EMBEDDING_MODEL,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

st.set_page_config(
    page_title="KSB Marker",
    page_icon="KSB",
    layout="wide",
    initial_sidebar_state="expanded",
)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# CSS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

st.markdown("""
<style>
    .main-header { font-size: 2.8rem; font-weight: 700; background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; text-align: center; margin-bottom: 0.5rem; }
    .sub-header { font-size: 1.1rem; color: #8b95a5; margin-bottom: 2rem; text-align: center; }
    .phase-card { background: linear-gradient(135deg, #1e2530 0%, #252d3a 100%); border-radius: 12px; padding: 1rem; border: 1px solid #3d4654; margin: 0.3rem 0; }
    .workspace-card { background: linear-gradient(160deg, #18212f 0%, #202d40 100%); border-radius: 14px; padding: 1rem; border: 1px solid #31435b; margin-bottom: 0.8rem; }
    .workspace-title { color: #e2e8f0; font-weight: 700; margin-bottom: 0.35rem; }
    .workspace-copy { color: #9fb0c5; font-size: 0.9rem; line-height: 1.5; }
    .report-document { padding: 1.35rem 1.45rem; max-height: 960px; overflow-y: auto; }
    .report-heading { color: #f4f7fb; font-family: Georgia, serif; font-weight: 700; line-height: 1.35; }
    .report-heading-0 { font-size: 1.75rem; margin: 0 0 1rem 0; }
    .report-heading-1 { font-size: 1.35rem; margin: 1.2rem 0 0.55rem 0; }
    .report-heading-2 { font-size: 1.16rem; margin: 1rem 0 0.45rem 0; }
    .report-heading-3 { font-size: 1.02rem; margin: 0.9rem 0 0.35rem 0; text-transform: none; }
    .report-paragraph { color: #dfe7f0; font-family: Georgia, serif; font-size: 1rem; line-height: 1.62; margin: 0 0 0.78rem 0; }
    .report-caption { color: #aebccd; font-family: Georgia, serif; font-size: 0.92rem; font-style: italic; line-height: 1.5; margin: 0.1rem 0 0.85rem 0; }
    .report-figure { margin: 1rem 0 1.1rem 0; text-align: center; }
    .report-figure img { max-width: 100%; height: auto; border-radius: 12px; border: 1px solid #31435b; background: #0f1723; }
    .report-table { margin: 0.9rem 0 1.2rem 0; overflow-x: auto; }
    .report-table table { width: 100%; border-collapse: collapse; font-size: 0.94rem; }
    .report-table th { background: #223048; color: #edf2f7; padding: 0.55rem 0.65rem; border: 1px solid #31435b; text-align: left; }
    .report-table td { color: #d7e0eb; padding: 0.55rem 0.65rem; border: 1px solid #31435b; text-align: left; vertical-align: top; }
    .phase-card.active { border-color: #667eea; box-shadow: 0 0 15px rgba(102, 126, 234, 0.3); }
    .phase-card.complete { border-color: #10b981; }
    .grade-badge { padding: 0.3rem 0.8rem; border-radius: 20px; font-weight: 600; display: inline-block; }
    .grade-merit { background: #3b82f6; color: white; }
    .grade-pass { background: #10b981; color: white; }
    .grade-referral { background: #ef4444; color: white; }
    .grade-unknown { background: #6b7280; color: white; }
    .stat-card { background: linear-gradient(135deg, #1e2530 0%, #252d3a 100%); border-radius: 10px; padding: 1rem; text-align: center; border: 1px solid #2d3748; }
    .stat-number { font-size: 2rem; font-weight: 700; }
    .tool-tag { background: #2d3748; padding: 0.15rem 0.4rem; border-radius: 4px; font-size: 0.65rem; margin: 0.1rem; display: inline-block; color: #a0aec0; }
</style>
""", unsafe_allow_html=True)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Session state
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def init_session_state():
    defaults = {
        "report_data": None,
        "ksb_criteria": None,
        "assignment_brief": None,
        "assignment_brief_source": "default",
        "assignment_brief_filename": None,
        "agent_results": None,
        "assessment_complete": False,
        "selected_module": "DSP",
        "verbose_mode": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Cached resources
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@st.cache_resource
def load_ollama_client():
    try:
        return OllamaClient(base_url=OLLAMA_BASE_URL, model=OLLAMA_MODEL, timeout=180)
    except Exception as e:
        logger.error(f"Failed to connect to Ollama: {e}")
        return None


@st.cache_resource
def load_embedder():
    try:
        return Embedder(model_name=EMBEDDING_MODEL)
    except Exception as e:
        logger.error(f"Failed to load embedder: {e}")
        return None


@st.cache_resource
def load_image_processor():
    try:
        return ImageProcessor(max_size=(1024, 1024))
    except Exception as e:
        return None


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Document processing
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def process_report(uploaded_file, image_processor=None) -> Optional[Dict[str, Any]]:
    """Process uploaded report and return structured data."""
    if uploaded_file is None:
        return None

    try:
        suffix = Path(uploaded_file.name).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name

        processor = DocxProcessor() if suffix == ".docx" else PDFProcessor()
        doc = processor.process(tmp_path)

        raw_text = getattr(doc, "raw_text", "")
        if len(raw_text.strip()) < 200:
            st.warning("Document appears too short to assess (less than 200 characters).")
            Path(tmp_path).unlink(missing_ok=True)
            return None

        chunker = SmartChunker()
        chunks = chunker.chunk_report(doc.chunks, document_id="report")

        if len(chunks) < 3:
            st.warning(f"Document produced very few text sections ({len(chunks)} chunks). Results may be unreliable.")

        images = []
        if image_processor:
            if suffix == ".docx" and hasattr(doc, "figures") and doc.figures:
                images = image_processor.process_docx_images(
                    doc.figures, getattr(doc, "figure_captions", {})
                )
            elif suffix == ".pdf":
                images = image_processor.process_pdf_images(tmp_path)

        total_pages = getattr(doc, "total_pages", None) or getattr(doc, "total_pages_estimate", 1)

        render_blocks = []
        if suffix == ".docx":
            render_blocks = _build_docx_render_blocks(
                tmp_path,
                getattr(doc, "figures", {}),
                images,
            )

        return {
            "chunks": [c.to_dict() if hasattr(c, "to_dict") else c for c in chunks],
            "title": doc.title or uploaded_file.name,
            "filename": uploaded_file.name,
            "file_extension": suffix,
            "file_bytes": uploaded_file.getvalue(),
            "total_pages": total_pages,
            "pages_are_accurate": getattr(doc, "pages_are_accurate", suffix == ".pdf"),
            "raw_text": raw_text,
            "images": images,
            "render_blocks": render_blocks,
            "tmp_path": tmp_path,
        }
    except Exception as e:
        logger.exception("Error processing report")
        st.error(f"Error: {str(e)}")
        return None


def process_assignment_brief(uploaded_file, module_code: str):
    """Process an uploaded assignment brief and return a parsed brief object."""
    if uploaded_file is None:
        return None

    suffix = Path(uploaded_file.name).suffix.lower()
    tmp_path = None

    try:
        if suffix == ".txt":
            raw_text = uploaded_file.getvalue().decode("utf-8", errors="replace")
        else:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(uploaded_file.getvalue())
                tmp_path = tmp.name

            if suffix == ".docx":
                raw_text = DocxProcessor().process(tmp_path).raw_text
            elif suffix == ".pdf":
                raw_text = PDFProcessor().process(tmp_path).raw_text
            else:
                st.error("Unsupported assignment brief format.")
                return None

        if len(raw_text.strip()) < 100:
            st.warning("Assignment brief appears too short to parse reliably.")
            return None

        return parse_uploaded_brief(raw_text, module_code)
    except Exception as e:
        logger.exception("Error processing assignment brief")
        st.error(f"Brief error: {str(e)}")
        return None
    finally:
        if tmp_path:
            Path(tmp_path).unlink(missing_ok=True)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# LangGraph pipeline
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def run_langgraph_pipeline(
    report_data: dict,
    ksb_criteria: list,
    assignment_brief,
    embedder: Embedder,
    module_code: str,
    progress_bar,
    status_text,
    phase_placeholder,
    verbose_log: Optional[list] = None,
) -> Optional[dict]:
    """
    Run the LangGraph assessment pipeline with streaming progress updates.

    Returns the final graph state dict.
    """
    from src.graph import build_graph

    tmpdir = tempfile.mkdtemp()
    vector_store = None

    try:
        # â”€â”€ Step 1: Index report into ChromaDB â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        progress_bar.progress(0.05, "Indexing report into vector store...")
        status_text.text("Indexing report...")
        _update_phase_display(phase_placeholder, "retrieval", active=False)

        vector_store = ChromaStore(
            persist_directory=tmpdir,
        )
        report_texts = [c.get("content", "") for c in report_data["chunks"]]
        report_embeddings = embedder.embed_documents(report_texts)
        vector_store.add_report(report_data["chunks"], report_embeddings)

        if verbose_log is not None:
            verbose_log.append(f"[INDEX] Indexed {len(report_data['chunks'])} chunks into ChromaDB")

        # â”€â”€ Step 2: Prepare initial state â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        ksb_list = []
        for ksb in ksb_criteria:
            if hasattr(ksb, "code"):
                ksb_list.append({
                    "code": ksb.code,
                    "title": ksb.title,
                    "full_description": ksb.full_description,
                    "pass_criteria": ksb.pass_criteria,
                    "merit_criteria": ksb.merit_criteria,
                    "referral_criteria": ksb.referral_criteria,
                    "category": ksb.category,
                })
            else:
                ksb_list.append(ksb)

        brief_dict = {}
        if assignment_brief:
            brief_dict = assignment_brief.to_dict() if hasattr(assignment_brief, "to_dict") else assignment_brief

        images = []
        for img in report_data.get("images", []):
            if hasattr(img, "image_id"):
                images.append({
                    "image_id": img.image_id,
                    "caption": getattr(img, "caption", ""),
                    "base64": getattr(img, "base64_data", ""),
                })
            elif isinstance(img, dict):
                images.append(img)

        initial_state = {
            "module_code": module_code,
            "report_chunks": report_data["chunks"],
            "report_images": images,
            "ksb_criteria": ksb_list,
            "assignment_brief": brief_dict,
            "pages_are_accurate": report_data.get("pages_are_accurate", False),
            "evidence_map": {},
            "content_quality": {},
            "image_analyses": [],
            "ksb_scores": {},
            "overall_recommendation": "",
            "content_warnings": [],
            "ksb_feedback": {},
            "overall_feedback": "",
            "current_ksb_index": 0,
            "errors": [],
            "phase": "retrieval",
        }

        # â”€â”€ Step 3: Build and stream graph â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        # Monkey-patch config so retriever_node can find the temp store
        import config
        original_persist = config.CHROMA_PERSIST_DIR
        config.CHROMA_PERSIST_DIR = tmpdir

        try:
            graph = build_graph()

            # Phase progress mapping
            phase_progress = {
                "retriever": (0.30, "Retrieving evidence..."),
                "DSP": (0.75, "Scoring KSBs (DSP)..."),
                "MLCC": (0.75, "Scoring KSBs (MLCC)..."),
                "AIDI": (0.75, "Scoring KSBs (AIDI)..."),
                "feedback": (0.95, "Generating feedback..."),
            }

            # Accumulate node outputs into a single state dict
            accumulated_state = dict(initial_state)

            for chunk in graph.stream(initial_state):
                for node_name, node_output in chunk.items():
                    # Merge node output into accumulated state
                    accumulated_state.update(node_output)

                    if node_name in phase_progress:
                        end_pct, label = phase_progress[node_name]
                        progress_bar.progress(end_pct, label)
                        status_text.text(label)

                        # Determine phase for display
                        if node_name == "retriever":
                            _update_phase_display(phase_placeholder, "retrieval")
                        elif node_name in ("DSP", "MLCC", "AIDI"):
                            _update_phase_display(phase_placeholder, "scoring")
                        elif node_name == "feedback":
                            _update_phase_display(phase_placeholder, "feedback")

                        if verbose_log is not None:
                            verbose_log.append(f"[GRAPH] Node '{node_name}' completed")
                            if node_name == "retriever":
                                em = node_output.get("evidence_map", {})
                                total_chunks = sum(
                                    e.get("total_retrieved", 0)
                                    for e in em.values()
                                )
                                verbose_log.append(
                                    f"[RETRIEVER] {len(em)} KSBs, {total_chunks} total evidence chunks"
                                )
                            elif node_name in ("DSP", "MLCC", "AIDI"):
                                scores = node_output.get("ksb_scores", {})
                                for code, s in scores.items():
                                    verbose_log.append(
                                        f"[SCORE] {code}: {s.get('grade', '?')} "
                                        f"(confidence={s.get('confidence', '?')}, "
                                        f"method={s.get('extraction_method', '?')})"
                                    )

            progress_bar.progress(1.0, "Assessment complete!")
            status_text.text("Assessment complete!")
            _update_phase_display(phase_placeholder, "complete")

            return accumulated_state

        finally:
            config.CHROMA_PERSIST_DIR = original_persist

    finally:
        # Clean up
        if vector_store is not None:
            vector_store.close()
            del vector_store
        import gc
        gc.collect()
        time.sleep(0.3)
        try:
            shutil.rmtree(tmpdir)
        except Exception as e:
            logger.warning(f"Failed to clean up temp directory: {e}")

        if report_data and "tmp_path" in report_data:
            Path(report_data["tmp_path"]).unlink(missing_ok=True)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Phase display
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def _update_phase_display(placeholder, current_phase: str, active: bool = True):
    """Render the three phase cards with current state."""
    if placeholder is None:
        return

    phases = [
        ("retrieval", "Retrieval", "Hybrid BM25+semantic evidence search"),
        ("scoring", "Scoring", "KSB evaluation via local LLM"),
        ("feedback", "Feedback", "Structured feedback generation"),
    ]

    cols = placeholder.columns(3)
    for i, (phase_key, label, desc) in enumerate(phases):
        if current_phase == "complete":
            css_class = "complete"
            icon = "Done"
        elif phase_key == current_phase and active:
            css_class = "active"
            icon = "Running..."
        elif phases.index((phase_key, label, desc)) < [p[0] for p in phases].index(current_phase) if current_phase in [p[0] for p in phases] else False:
            css_class = "complete"
            icon = "Done"
        else:
            css_class = ""
            icon = "Waiting"

        with cols[i]:
            st.markdown(f"""
            <div class="phase-card {css_class}">
                <div style="font-weight: 600; color: #e2e8f0;">{label}</div>
                <div style="font-size: 0.8rem; color: #8b95a5;">{desc}</div>
                <div style="font-size: 0.75rem; color: #667eea; margin-top: 0.3rem;">{icon}</div>
            </div>
            """, unsafe_allow_html=True)


def _build_docx_render_blocks(docx_path: str, figures: Dict[str, bytes], processed_images: list) -> list:
    """Build ordered render blocks for DOCX viewing with inline images."""
    try:
        document = Document(docx_path)
    except Exception:
        logger.warning("Could not build DOCX render blocks", exc_info=True)
        return []

    paragraph_lookup = {para._element: para for para in document.paragraphs}
    table_lookup = {table._tbl: table for table in document.tables}
    figure_id_by_hash = {
        hashlib.sha1(raw_bytes).hexdigest(): figure_id
        for figure_id, raw_bytes in (figures or {}).items()
    }
    processed_images_by_id = {
        image.image_id: image
        for image in (processed_images or [])
        if getattr(image, "image_id", None)
    }
    processor = DocxProcessor()
    blocks = []

    for element in document.element.body:
        if element.tag.endswith('p'):
            para = paragraph_lookup.get(element)
            if para is None:
                continue

            style_name = para.style.name if para.style else ""
            heading_level = processor._detect_heading_level(style_name, para)
            text = para.text.strip()

            if text:
                block_type = "heading" if heading_level is not None else "paragraph"
                if block_type == "paragraph" and "caption" in style_name.lower():
                    block_type = "caption"

                blocks.append({
                    "type": block_type,
                    "text": text,
                    "level": heading_level if heading_level is not None else 1,
                })

            for node in para._element.iter():
                if not str(node.tag).endswith('}blip'):
                    continue

                rel_id = node.get(qn('r:embed'))
                if not rel_id or rel_id not in document.part.rels:
                    continue

                try:
                    raw_bytes = document.part.rels[rel_id].target_part.blob
                except Exception:
                    continue

                figure_id = figure_id_by_hash.get(hashlib.sha1(raw_bytes).hexdigest())
                processed_image = processed_images_by_id.get(figure_id)
                if not processed_image or not getattr(processed_image, "base64_data", None):
                    continue

                blocks.append({
                    "type": "image",
                    "base64_data": processed_image.base64_data,
                    "format": processed_image.format or "png",
                    "image_id": processed_image.image_id,
                })

        elif element.tag.endswith('tbl'):
            table = table_lookup.get(element)
            if table is None or not table.rows:
                continue

            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])

            if rows:
                blocks.append({
                    "type": "table",
                    "rows": rows,
                })

    return blocks


def _render_report_text_preview(report_data: dict):
    """Render a readable text preview for DOCX/PDF reports."""
    raw_text = (report_data or {}).get("raw_text", "").strip()
    if not raw_text:
        st.info("No readable text preview is available for this report.")
        return

    text_preview = raw_text if len(raw_text) <= 120000 else raw_text[:120000] + "\n\n[Preview truncated for display]"
    st.markdown(
        f"""
        <div class="workspace-card" style="max-height: 920px; overflow-y: auto; white-space: pre-wrap; font-family: Georgia, serif; line-height: 1.45;">
            <div class="workspace-title">Reading View</div>
            <div class="workspace-copy" style="color: #d7e0eb;">{html.escape(text_preview)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _build_docx_report_html(report_data: dict) -> Optional[str]:
    """Build an HTML view for DOCX reports with inline images and controlled spacing."""
    blocks = (report_data or {}).get("render_blocks", []) or []
    if not blocks:
        return None

    parts = ['<div class="workspace-card report-document">']

    for block in blocks:
        block_type = block.get("type")

        if block_type in {"heading", "paragraph", "caption"}:
            text_html = html.escape(block.get("text", "")).replace("\n", "<br>")
            if not text_html:
                continue

            if block_type == "heading":
                level = max(0, min(int(block.get("level", 1)), 3))
                parts.append(f'<div class="report-heading report-heading-{level}">{text_html}</div>')
            elif block_type == "caption":
                parts.append(f'<p class="report-caption">{text_html}</p>')
            else:
                parts.append(f'<p class="report-paragraph">{text_html}</p>')

        elif block_type == "image":
            base64_data = block.get("base64_data")
            if not base64_data:
                continue
            img_format = html.escape(block.get("format", "png"))
            image_id = html.escape(block.get("image_id", "report image"))
            parts.append(
                f'<figure class="report-figure"><img src="data:image/{img_format};base64,{base64_data}" alt="{image_id}"></figure>'
            )

        elif block_type == "table":
            rows = block.get("rows", [])
            if not rows:
                continue

            header_cells = ''.join(f'<th>{html.escape(str(cell))}</th>' for cell in rows[0])
            body_rows = []
            for row in rows[1:]:
                body_cells = ''.join(f'<td>{html.escape(str(cell))}</td>' for cell in row)
                body_rows.append(f'<tr>{body_cells}</tr>')

            parts.append(
                '<div class="report-table"><table>'
                f'<thead><tr>{header_cells}</tr></thead>'
                f'<tbody>{"".join(body_rows)}</tbody>'
                '</table></div>'
            )

    parts.append('</div>')
    return ''.join(parts)


def _render_docx_report_view(report_data: dict):
    """Render a DOCX report view that keeps paragraph flow and images inline."""
    report_html = _build_docx_report_html(report_data)
    if not report_html:
        _render_report_text_preview(report_data)
        return

    st.markdown(report_html, unsafe_allow_html=True)


def display_report_viewer(module_code: str, module_name: str, llm, embedder, report_data, assignment_brief, results: Optional[dict] = None):
    """Render a full-width live report viewer so users can read the submission while assessment runs."""
    st.markdown("## Report Viewer")

    if hasattr(assignment_brief, "tasks"):
        brief_tasks = len(assignment_brief.tasks)
    elif isinstance(assignment_brief, dict):
        brief_tasks = len(assignment_brief.get("tasks", []))
    else:
        brief_tasks = 0

    report_ready = bool(report_data)
    llm_ready = llm is not None
    enough_chunks = len(report_data.get("chunks", [])) >= 3 if report_data else False

    status_bits = [f"Module: {module_code}", f"Brief tasks: {brief_tasks}"]
    if report_ready:
        status_bits.append(f"Pages: {report_data.get('total_pages', 0)}")
        status_bits.append(f"Chunks: {len(report_data.get('chunks', []))}")
    st.caption(" | ".join(status_bits))

    if not llm_ready:
        st.warning("Ollama is offline. You can still read the report here while the assessment setup is being resolved.")
    elif report_ready and not enough_chunks:
        st.warning("This report loaded, but it produced very few chunks, so the assessment may be unreliable.")
    elif report_ready and results and results.get("ksb_scores"):
        st.success(f"Latest overall recommendation: {results.get('overall_recommendation', 'UNKNOWN')}")
    elif report_ready:
        st.info("The report is ready to read. Start the assessment when you are ready.")
    else:
        st.info("Upload a report to open a live reading view here.")
        return

    extension = report_data.get("file_extension", "").lower()
    file_bytes = report_data.get("file_bytes", b"")

    if extension == ".pdf" and file_bytes:
        pdf_tab, text_tab = st.tabs(["PDF", "Reading View"])

        with pdf_tab:
            pdf_base64 = base64.b64encode(file_bytes).decode("utf-8")
            st.markdown(
                f"""
                <div class="workspace-card" style="padding: 0.4rem;">
                    <iframe src="data:application/pdf;base64,{pdf_base64}" width="100%" height="920" style="border: none; border-radius: 10px;"></iframe>
                </div>
                """,
                unsafe_allow_html=True,
            )
        with text_tab:
            _render_report_text_preview(report_data)
    elif extension == ".docx":
        doc_tab, text_tab = st.tabs(["Document View", "Reading View"])
        with doc_tab:
            st.caption("DOCX files are shown in a document view with inline extracted images.")
            _render_docx_report_view(report_data)
        with text_tab:
            _render_report_text_preview(report_data)
    else:
        _render_report_text_preview(report_data)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Content quality banners
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def display_content_quality_banners(content_quality: dict, content_warnings: list):
    """Display content quality banners based on detection results."""
    if not content_quality:
        return

    status = content_quality.get("status", "OK")
    adversarial_ksbs = content_quality.get("adversarial_ksbs", [])
    off_topic_count = content_quality.get("off_topic_count", 0)

    if adversarial_ksbs and off_topic_count > 0:
        st.error(
            f"**Adversarial Content Detected**\n\n"
            f"**{len(adversarial_ksbs)}** adversarial reflection table(s) found "
            f"affecting KSBs: {', '.join(adversarial_ksbs)}. "
            f"Additionally, **{off_topic_count}** off-topic section(s) detected. "
            f"Affected KSBs have been automatically referred."
        )
    elif adversarial_ksbs:
        st.error(
            f"**Adversarial Content Detected**\n\n"
            f"**{len(adversarial_ksbs)}** adversarial reflection table(s) found "
            f"affecting KSBs: {', '.join(adversarial_ksbs)}. "
            f"These KSBs have been automatically referred."
        )
    elif off_topic_count > 0:
        st.warning(
            f"**Content Quality Warning**\n\n"
            f"**{off_topic_count}** section(s) contain content with very low relevance. "
            f"Manual review recommended."
        )
    elif status == "CRITICAL":
        st.warning("Content quality issues detected. Review results carefully.")
    elif status == "WARNING":
        st.info("Some sections contain content with low relevance to the module.")

    # Show specific content warnings from specialist
    for warning in content_warnings:
        st.warning(warning)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Results display
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def display_results(results: dict):
    """Display assessment results from LangGraph state."""

    ksb_scores = results.get("ksb_scores", {})
    ksb_feedback = results.get("ksb_feedback", {})
    overall_recommendation = results.get("overall_recommendation", "UNKNOWN")
    content_quality = results.get("content_quality", {})
    content_warnings = results.get("content_warnings", [])
    ksb_criteria = results.get("ksb_criteria", [])
    errors = results.get("errors", [])

    st.markdown("## Assessment Results")

    # Content quality banners
    display_content_quality_banners(content_quality, content_warnings)

    # Errors
    if errors:
        with st.expander(f"Errors ({len(errors)})", expanded=False):
            for err in errors:
                st.error(err)

    # Grade distribution
    grades = [s["grade"] for s in ksb_scores.values()]
    merit_count = grades.count("MERIT")
    pass_count = grades.count("PASS")
    referral_count = grades.count("REFERRAL")
    total = len(grades)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(f'<div class="stat-card"><div class="stat-number" style="color:#fff">{total}</div><div>Total KSBs</div></div>', unsafe_allow_html=True)
    with col2:
        st.markdown(f'<div class="stat-card"><div class="stat-number" style="color:#3b82f6">{merit_count}</div><div>Merit</div></div>', unsafe_allow_html=True)
    with col3:
        st.markdown(f'<div class="stat-card"><div class="stat-number" style="color:#10b981">{pass_count}</div><div>Pass</div></div>', unsafe_allow_html=True)
    with col4:
        st.markdown(f'<div class="stat-card"><div class="stat-number" style="color:#ef4444">{referral_count}</div><div>Referral</div></div>', unsafe_allow_html=True)

    # Overall recommendation badge
    grade_class = f"grade-{overall_recommendation.lower()}" if overall_recommendation in ("MERIT", "PASS", "REFERRAL") else "grade-unknown"
    st.markdown(f"""
    <div style="text-align: center; margin: 1rem 0;">
        <span class="grade-badge {grade_class}">{overall_recommendation}</span>
        <span style="color: #8b95a5; margin-left: 0.5rem;">Overall Recommendation</span>
    </div>
    """, unsafe_allow_html=True)

    # Export buttons
    st.markdown("---")
    st.markdown("### Export Results")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    module = results.get("module_code", "UNKNOWN")

    col1, col2, col3 = st.columns(3)

    with col1:
        csv_data = _export_csv(ksb_scores, ksb_feedback, overall_recommendation, module)
        st.download_button(
            label="Download CSV",
            data=csv_data,
            file_name=f"ksb_assessment_{module}_{timestamp}.csv",
            mime="text/csv",
            use_container_width=True,
        )

    with col2:
        json_data = json.dumps({
            "module_code": module,
            "overall_recommendation": overall_recommendation,
            "ksb_scores": {k: dict(v) for k, v in ksb_scores.items()},
            "ksb_feedback": {k: dict(v) for k, v in ksb_feedback.items()},
            "content_quality": content_quality,
            "errors": errors,
        }, indent=2, default=str)
        st.download_button(
            label="Download JSON",
            data=json_data,
            file_name=f"ksb_assessment_{module}_{timestamp}.json",
            mime="application/json",
            use_container_width=True,
        )

    with col3:
        md_parts = [results.get("overall_feedback", "")]
        for code in sorted(ksb_feedback.keys()):
            fb = ksb_feedback[code]
            md_parts.append(fb.get("formatted_markdown", ""))
        md_data = "\n\n---\n\n".join(md_parts)
        st.download_button(
            label="Download Markdown",
            data=md_data,
            file_name=f"feedback_{module}_{timestamp}.md",
            mime="text/markdown",
            use_container_width=True,
        )

    st.markdown("---")

    # Overall feedback (strip the grade table header to avoid duplication with stat cards above)
    if results.get("overall_feedback"):
        with st.expander("Overall Assessment Summary", expanded=True):
            overall_text = results["overall_feedback"]
            # Remove the auto-generated grade table header (everything before the LLM response)
            separator = "---\n\n"
            if separator in overall_text:
                # The header ends at the last "---" separator before the LLM content
                parts = overall_text.split(separator)
                # The LLM-generated content is after the last separator
                llm_content = parts[-1].strip() if len(parts) > 1 else overall_text
                st.markdown(llm_content)
            else:
                st.markdown(overall_text)

    # Per-KSB breakdown
    st.markdown("## KSB Breakdown")

    # Build criteria lookup for display
    criteria_lookup = {c["code"]: c for c in ksb_criteria} if ksb_criteria else {}

    # Use criteria order (matches pipeline processing order), fall back to sorted
    ksb_display_order = [c["code"] for c in ksb_criteria if c["code"] in ksb_scores] if ksb_criteria else sorted(ksb_scores.keys())

    for ksb_code in ksb_display_order:
        score = ksb_scores[ksb_code]
        grade = score["grade"]
        confidence = score.get("confidence", "")
        criterion = criteria_lookup.get(ksb_code, {})
        ksb_title = criterion.get("title", score.get("ksb_code", ksb_code))

        with st.expander(f"**{ksb_code}** - {ksb_title} [{grade}]"):
            # Feedback
            fb = ksb_feedback.get(ksb_code, {})
            if fb.get("formatted_markdown"):
                st.markdown(fb["formatted_markdown"])

            # Transparency tabs
            tab1, tab2, tab3, tab4 = st.tabs([
                "Evidence", "LLM Reasoning", "Validation", "Grade Decision"
            ])

            with tab1:
                evidence_map = results.get("evidence_map", {})
                evidence = evidence_map.get(ksb_code, {})
                chunks = evidence.get("chunks", [])
                st.caption(
                    f"**Retrieved:** {evidence.get('total_retrieved', 0)} chunks | "
                    f"**Strategy:** {evidence.get('search_strategy', 'N/A')} | "
                    f"**Query variations:** {len(evidence.get('query_variations', []))}"
                )
                if chunks:
                    for idx, chunk in enumerate(chunks[:5], 1):
                        metadata = chunk.get("metadata", {})
                        section = metadata.get("section_number", "") or metadata.get("section_title", "")
                        sim = chunk.get("similarity", 0)
                        st.caption(f"**Chunk {idx}** | Section: `{section or 'N/A'}` | Relevance: `{sim:.3f}`")
                        content = chunk.get("content", "")
                        display = content[:500] + "..." if len(content) > 500 else content
                        st.markdown(
                            f"<div style='background-color: rgba(255,255,255,0.05); "
                            f"padding: 10px; border-radius: 5px; border-left: 3px solid #4CAF50; "
                            f"margin-bottom: 0.5rem; font-family: monospace; font-size: 0.85em;'>"
                            f"{display}</div>",
                            unsafe_allow_html=True,
                        )
                else:
                    st.info("No evidence chunks retrieved for this KSB.")

            with tab2:
                raw_response = score.get("raw_llm_response", "")
                if raw_response:
                    st.code(raw_response, language="text")
                else:
                    st.info("No LLM response (auto-graded).")

            with tab3:
                audit = score.get("audit_trail", {})
                action = audit.get("validation_action", "N/A")
                conf = audit.get("validation_confidence", 0.0)
                warnings = audit.get("validation_warnings", [])
                val_errors = audit.get("validation_errors", [])

                action_color = "#10b981" if action == "accept" else "#f59e0b" if action == "flag_for_review" else "#ef4444"
                st.markdown(
                    f"<span style='background-color: {action_color}; color: white; "
                    f"padding: 0.3rem 0.8rem; border-radius: 20px; font-weight: 600;'>"
                    f"{action.upper() if isinstance(action, str) else action}</span>",
                    unsafe_allow_html=True,
                )
                st.metric("Validation Confidence", f"{conf:.2f}" if isinstance(conf, (int, float)) else str(conf))

                if val_errors:
                    st.markdown("**Errors:**")
                    for err in val_errors:
                        st.error(err)
                if warnings:
                    st.markdown("**Warnings:**")
                    for w in warnings:
                        st.warning(w)
                if not val_errors and not warnings:
                    st.success("No validation issues.")

            with tab4:
                grade_color = "#3b82f6" if grade == "MERIT" else "#10b981" if grade == "PASS" else "#ef4444"
                st.markdown(
                    f"<div style='text-align: center; margin: 1rem 0;'>"
                    f"<span style='background-color: {grade_color}; color: white; "
                    f"padding: 0.5rem 1.5rem; border-radius: 20px; font-weight: 700; "
                    f"font-size: 1.2rem;'>{grade}</span></div>",
                    unsafe_allow_html=True,
                )

                c1, c2 = st.columns(2)
                with c1:
                    st.markdown("**Criteria Met:**")
                    st.checkbox("Pass Criteria", value=score.get("pass_criteria_met", False), disabled=True, key=f"{ksb_code}_pass")
                    st.checkbox("Merit Criteria", value=score.get("merit_criteria_met", False), disabled=True, key=f"{ksb_code}_merit")
                with c2:
                    st.markdown("**Assessment Metrics:**")
                    st.metric("Evidence Strength", score.get("evidence_strength", "N/A"))
                    st.metric("Confidence", confidence)

                st.caption(f"**Extraction Method:** {score.get('extraction_method', 'N/A')}")
                if score.get("placeholder_detected"):
                    st.warning("Placeholder content detected in evidence.")
                if score.get("adversarial_detected"):
                    st.error("Adversarial content detected for this KSB.")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# Export helpers
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def _export_csv(
    ksb_scores: dict, ksb_feedback: dict,
    overall_recommendation: str, module_code: str,
) -> str:
    """Generate CSV export from graph results."""
    output = StringIO()
    writer = csv.writer(output)

    writer.writerow([
        "KSB Code", "Grade", "Confidence", "Evidence Strength",
        "Pass Met", "Merit Met", "Extraction Method",
        "Strengths", "Improvements", "Rationale",
    ])

    for code in sorted(ksb_scores.keys()):
        s = ksb_scores[code]
        fb = ksb_feedback.get(code, {})
        strengths = _normalize_feedback_items(fb.get("strengths", []))
        improvements = _normalize_feedback_items(fb.get("improvements", []))
        writer.writerow([
            code,
            s.get("grade", ""),
            s.get("confidence", ""),
            s.get("evidence_strength", ""),
            s.get("pass_criteria_met", ""),
            s.get("merit_criteria_met", ""),
            s.get("extraction_method", ""),
            " | ".join(strengths),
            " | ".join(improvements),
            s.get("rationale", "")[:200],
        ])

    grades = [s["grade"] for s in ksb_scores.values()]
    writer.writerow([])
    writer.writerow(["SUMMARY"])
    writer.writerow(["Module", module_code])
    writer.writerow(["Overall", overall_recommendation])
    writer.writerow(["Merit", grades.count("MERIT")])
    writer.writerow(["Pass", grades.count("PASS")])
    writer.writerow(["Referral", grades.count("REFERRAL")])

    return output.getvalue()


def _normalize_feedback_items(items) -> list[str]:
    """Convert mixed feedback item formats into displayable strings."""
    normalized = []
    for item in items or []:
        if isinstance(item, dict):
            text = (
                item.get("suggestion")
                or item.get("strength")
                or item.get("area")
                or item.get("evidence")
                or str(item)
            )
        else:
            text = str(item)

        text = text.strip()
        if text:
            normalized.append(text)

    return normalized


def export_results_to_csv(results: dict, module_code: str = "") -> str:
    """Backward-compatible CSV export helper used by local verification scripts."""
    if not results:
        return ""

    if "ksb_scores" in results:
        resolved_module = module_code or results.get("module_code", "UNKNOWN")
        return _export_csv(
            results.get("ksb_scores", {}),
            results.get("ksb_feedback", {}),
            results.get("overall_recommendation", "UNKNOWN"),
            resolved_module,
        )

    scoring_results = results.get("scoring_results", [])
    feedback_results = results.get("feedback_results", [])
    overall_summary = results.get("overall_summary", {})

    ksb_scores = {}
    for item in scoring_results:
        code = item.get("ksb_code")
        if not code:
            continue
        ksb_scores[code] = {
            "grade": item.get("grade", ""),
            "confidence": item.get("confidence", ""),
            "evidence_strength": item.get("evidence_strength", ""),
            "pass_criteria_met": item.get("pass_criteria_met", item.get("pass_met", "")),
            "merit_criteria_met": item.get("merit_criteria_met", item.get("merit_met", "")),
            "extraction_method": item.get("extraction_method", ""),
            "rationale": item.get("rationale", ""),
        }

    ksb_feedback = {}
    for item in feedback_results:
        code = item.get("ksb_code")
        if not code:
            continue
        ksb_feedback[code] = {
            "strengths": _normalize_feedback_items(item.get("strengths", [])),
            "improvements": _normalize_feedback_items(item.get("improvements", [])),
            "formatted_markdown": item.get("formatted_markdown") or item.get("formatted_feedback", ""),
        }

    resolved_module = (
        module_code
        or results.get("module_code")
        or overall_summary.get("module_code")
        or "UNKNOWN"
    )

    csv_output = _export_csv(
        ksb_scores,
        ksb_feedback,
        overall_summary.get("overall_recommendation", "UNKNOWN"),
        resolved_module,
    )

    if not overall_summary:
        return csv_output

    output = StringIO()
    output.write(csv_output)
    writer = csv.writer(output)

    writer.writerow([])
    writer.writerow(["OVERALL SUMMARY"])
    writer.writerow(["Total KSBs", overall_summary.get("total_ksbs", len(ksb_scores))])
    writer.writerow(["Merit Count", overall_summary.get("merit_count", 0)])
    writer.writerow(["Pass Count", overall_summary.get("pass_count", 0)])
    writer.writerow(["Referral Count", overall_summary.get("referral_count", 0)])
    writer.writerow(["Overall Recommendation", overall_summary.get("overall_recommendation", "UNKNOWN")])
    writer.writerow(["Confidence", overall_summary.get("confidence", "")])

    key_strengths = _normalize_feedback_items(overall_summary.get("key_strengths", []))
    if key_strengths:
        writer.writerow([])
        writer.writerow(["KEY STRENGTHS"])
        for item in key_strengths:
            writer.writerow([item])

    priority_improvements = _normalize_feedback_items(overall_summary.get("priority_improvements", []))
    if priority_improvements:
        writer.writerow([])
        writer.writerow(["PRIORITY IMPROVEMENTS"])
        for item in priority_improvements:
            writer.writerow([item])

    return output.getvalue()

def main():
    init_session_state()

    st.markdown('<p class="main-header">KSB Coursework Marker</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Evidence-Led Coursework Review</p>', unsafe_allow_html=True)

    # â”€â”€ Sidebar â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    with st.sidebar:
        st.markdown("## Status")

        llm = load_ollama_client()
        embedder = load_embedder()
        image_processor = load_image_processor()

        if llm:
            st.success(f"Ollama: {OLLAMA_MODEL}")
        else:
            st.error("Ollama not connected")

        if embedder:
            st.success(f"Embedder ({embedder.embedding_dim}d)")
        else:
            st.error("Embedder error")

        if image_processor:
            st.success("Vision ready")

        st.markdown("---")

        # Module selection
        st.markdown("## Module")
        modules = get_available_modules()
        module_options = {code: info["name"] for code, info in modules.items()}
        selected = st.selectbox(
            "Module",
            list(module_options.keys()),
            format_func=lambda x: module_options[x],
            index=list(module_options.keys()).index(st.session_state.selected_module),
        )

        if selected != st.session_state.selected_module:
            st.session_state.selected_module = selected
            st.session_state.ksb_criteria = None
            st.session_state.assignment_brief = None
            st.session_state.assignment_brief_source = "default"
            st.session_state.assignment_brief_filename = None
            st.rerun()

        uploaded_brief = st.file_uploader(
            "Assignment brief (optional)",
            type=["docx", "pdf", "txt"],
            key=f"assignment_brief_{selected}",
            help=(
                "Upload the actual assignment brief for this module. "
                "If omitted, the built-in default brief will be used."
            ),
        )

        if uploaded_brief:
            brief_changed = (
                st.session_state.assignment_brief_source != "upload"
                or st.session_state.assignment_brief_filename != uploaded_brief.name
            )
            if brief_changed:
                with st.spinner("Parsing assignment brief..."):
                    parsed_brief = process_assignment_brief(uploaded_brief, selected)
                    if parsed_brief:
                        st.session_state.assignment_brief = parsed_brief
                        st.session_state.assignment_brief_source = "upload"
                        st.session_state.assignment_brief_filename = uploaded_brief.name
                        st.session_state.assessment_complete = False
                        st.session_state.agent_results = None
        elif st.session_state.assignment_brief_source == "upload":
            st.session_state.assignment_brief = None
            st.session_state.assignment_brief_source = "default"
            st.session_state.assignment_brief_filename = None
            st.session_state.assessment_complete = False
            st.session_state.agent_results = None

        if st.session_state.ksb_criteria is None:
            st.session_state.ksb_criteria = get_module_criteria(selected)

        if st.session_state.assignment_brief is None:
            st.session_state.assignment_brief = get_default_brief(selected)

        brief_obj = st.session_state.assignment_brief
        if hasattr(brief_obj, "tasks"):
            brief_tasks = len(brief_obj.tasks)
        elif isinstance(brief_obj, dict):
            brief_tasks = len(brief_obj.get("tasks", []))
        else:
            brief_tasks = 0

        brief_label = st.session_state.assignment_brief_filename or "Built-in default"
        st.caption(f"{len(st.session_state.ksb_criteria)} KSBs | Brief: {brief_tasks} tasks")
        st.caption(f"Brief source: {st.session_state.assignment_brief_source} ({brief_label})")

        st.markdown("---")
        st.session_state.verbose_mode = st.checkbox(
            "Verbose mode", value=st.session_state.verbose_mode
        )

        if st.button("Reset Session"):
            for key in ["report_data", "agent_results", "assessment_complete"]:
                st.session_state[key] = None if key != "assessment_complete" else False
            st.rerun()

    # â”€â”€ Main content â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    st.markdown("## Upload Report")
    uploaded_file = st.file_uploader("Student report", type=["docx", "pdf"])

    if uploaded_file:
        if (st.session_state.report_data is None or
                st.session_state.report_data.get("filename") != uploaded_file.name):
            with st.spinner("Processing document..."):
                report_data = process_report(uploaded_file, image_processor)
                if report_data:
                    st.session_state.report_data = report_data
                    st.session_state.assessment_complete = False
                    st.session_state.agent_results = None

        if st.session_state.report_data:
            r = st.session_state.report_data
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Title", r["title"][:30])
            c2.metric("Pages", r["total_pages"])
            c3.metric("Chunks", len(r["chunks"]))
            c4.metric("Images", len(r.get("images", [])))

    display_report_viewer(
        module_code=selected,
        module_name=module_options[selected],
        llm=llm,
        embedder=embedder,
        report_data=st.session_state.report_data,
        assignment_brief=st.session_state.assignment_brief,
        results=st.session_state.agent_results,
    )
    phase_placeholder = None

    # â”€â”€ Run button â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    st.markdown("---")

    can_run = bool(
        st.session_state.report_data
        and st.session_state.ksb_criteria
        and llm
        and embedder
        and len(st.session_state.report_data.get("chunks", [])) >= 3
    )

    if not can_run and st.session_state.report_data:
        chunk_count = len(st.session_state.report_data.get("chunks", []))
        if chunk_count < 3:
            st.error(f"Document has too few sections ({chunk_count} chunks). Cannot assess reliably.")

    if st.button(
        "Run Assessment",
        type="primary",
        disabled=not can_run,
        use_container_width=True,
    ):
        progress_bar = st.progress(0, "Starting...")
        status_text = st.empty()

        verbose_log = [] if st.session_state.verbose_mode else None
        verbose_container = None
        verbose_text = None
        if st.session_state.verbose_mode:
            verbose_container = st.expander("Verbose Log", expanded=True)
            verbose_text = verbose_container.empty()

        try:
            results = run_langgraph_pipeline(
                report_data=st.session_state.report_data,
                ksb_criteria=st.session_state.ksb_criteria,
                assignment_brief=st.session_state.assignment_brief,
                embedder=embedder,
                module_code=st.session_state.selected_module,
                progress_bar=progress_bar,
                status_text=status_text,
                phase_placeholder=phase_placeholder,
                verbose_log=verbose_log,
            )

            if results:
                st.session_state.agent_results = results
                st.session_state.assessment_complete = True

                if verbose_log and verbose_text:
                    verbose_text.code("\n".join(verbose_log), language="text")

                st.rerun()

        except Exception as e:
            logger.exception("Pipeline error")
            st.error(f"Error: {str(e)}")

    # â”€â”€ Display results â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if st.session_state.assessment_complete and st.session_state.agent_results:
        st.markdown("---")
        display_results(st.session_state.agent_results)


if __name__ == "__main__":
    main()











