"""
DOCX Processor - Extracts structured content from Word documents.
"""
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
import logging

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of document content with metadata."""
    content: str
    chunk_type: str
    heading_level: Optional[int] = None
    section_title: Optional[str] = None
    section_number: Optional[str] = None
    page_estimate: int = 1
    has_figure_reference: bool = False
    figure_ids: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ExtractedDocument:
    """Complete extracted document with all content and metadata."""
    chunks: List[DocumentChunk]
    title: Optional[str] = None
    total_pages_estimate: int = 1
    figures: Dict[str, bytes] = field(default_factory=dict)
    figure_captions: Dict[str, str] = field(default_factory=dict)
    raw_text: str = ""
    source_type: str = "docx"  # 'docx' or 'pdf' - affects page number reliability
    pages_are_accurate: bool = False  # True only for PDF


class DocxProcessor:
    """Processes DOCX files to extract structured content."""
    
    # More conservative estimate - typical Word doc with normal margins/font
    # is around 1800-2000 chars per page
    CHARS_PER_PAGE = 1800
    
    FIGURE_PATTERNS = [
        r'[Ff]igure\s+(\d+)',
        r'[Ff]ig\.?\s*(\d+)',
        r'[Ss]ee\s+[Ff]igure\s+(\d+)',
        r'[Ss]ee\s+[Ff]ig\.?\s*(\d+)',
        r'\([Ff]igure\s+(\d+)\)',
        r'\([Ff]ig\.?\s*(\d+)\)',
    ]
    
    def __init__(self):
        self.figure_pattern = re.compile('|'.join(self.FIGURE_PATTERNS))
    
    def process(self, file_path: str) -> ExtractedDocument:
        """Process a DOCX file and extract all content."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not path.suffix.lower() == '.docx':
            raise ValueError(f"Expected .docx file, got: {path.suffix}")
        
        doc = Document(file_path)
        
        chunks = []
        current_section = None
        current_section_number = None
        char_count = 0
        raw_text_parts = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = self._find_paragraph(doc, element)
                if para is not None:
                    chunk = self._process_paragraph(
                        para, current_section, current_section_number, char_count
                    )
                    if chunk:
                        if chunk.chunk_type == 'heading':
                            current_section = chunk.content
                            current_section_number = chunk.section_number
                        chunks.append(chunk)
                        raw_text_parts.append(chunk.content)
                        char_count += len(chunk.content)
            
            elif element.tag.endswith('tbl'):
                table = self._find_table(doc, element)
                if table is not None:
                    chunk = self._process_table(
                        table, current_section, current_section_number, char_count
                    )
                    if chunk:
                        chunks.append(chunk)
                        raw_text_parts.append(chunk.content)
                        char_count += len(chunk.content)
        
        figures, captions = self._extract_figures(doc)
        total_pages = max(1, char_count // self.CHARS_PER_PAGE + 1)
        title = self._extract_title(doc, chunks)
        
        return ExtractedDocument(
            chunks=chunks,
            title=title,
            total_pages_estimate=total_pages,
            figures=figures,
            figure_captions=captions,
            raw_text="\n\n".join(raw_text_parts),
            source_type="docx",
            pages_are_accurate=False  # DOCX pages are estimates only
        )
    
    def _find_paragraph(self, doc: Document, element) -> Optional[Any]:
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None
    
    def _find_table(self, doc: Document, element) -> Optional[Table]:
        for table in doc.tables:
            if table._tbl is element:
                return table
        return None
    
    def _process_paragraph(
        self, para, current_section: Optional[str],
        current_section_number: Optional[str], char_count: int
    ) -> Optional[DocumentChunk]:
        text = para.text.strip()
        if not text:
            return None
        
        style_name = para.style.name if para.style else ""
        heading_level = self._detect_heading_level(style_name, para)
        
        section_number = None
        if heading_level:
            section_number = self._extract_section_number(text)
        
        figure_refs = self.figure_pattern.findall(text)
        has_figure = len(figure_refs) > 0
        figure_ids = [f"figure_{ref}" for ref in figure_refs]
        
        page_estimate = max(1, char_count // self.CHARS_PER_PAGE + 1)
        
        if heading_level:
            return DocumentChunk(
                content=text,
                chunk_type='heading',
                heading_level=heading_level,
                section_title=text,
                section_number=section_number,
                page_estimate=page_estimate,
                has_figure_reference=has_figure,
                figure_ids=figure_ids,
                metadata={'style': style_name}
            )
        else:
            return DocumentChunk(
                content=text,
                chunk_type='paragraph',
                section_title=current_section,
                section_number=current_section_number,
                page_estimate=page_estimate,
                has_figure_reference=has_figure,
                figure_ids=figure_ids,
                metadata={'style': style_name}
            )
    
    def _extract_section_number(self, heading_text: str) -> Optional[str]:
        patterns = [
            r'^(\d+(?:\.\d+)*)\.\s',
            r'^(\d+(?:\.\d+)*)\s',
            r'^Part\s+(\d+)',
            r'^Section\s+(\d+)',
            r'^\[?(\d+(?:\.\d+)*)\]?\s*[-:)]',
        ]
        
        for pattern in patterns:
            match = re.match(pattern, heading_text, re.IGNORECASE)
            if match:
                return match.group(1)
        return None
    
    def _detect_heading_level(self, style_name: str, para) -> Optional[int]:
        style_lower = style_name.lower()
        
        if 'heading' in style_lower:
            match = re.search(r'(\d+)', style_name)
            if match:
                return int(match.group(1))
            return 1
        
        if 'title' in style_lower:
            return 0
        
        if para._element.pPr is not None:
            outline_lvl = para._element.pPr.find(qn('w:outlineLvl'))
            if outline_lvl is not None:
                val = outline_lvl.get(qn('w:val'))
                if val is not None:
                    return int(val) + 1
        
        return None
    
    def _process_table(
        self, table: Table, current_section: Optional[str],
        current_section_number: Optional[str], char_count: int
    ) -> Optional[DocumentChunk]:
        if not table.rows:
            return None
        
        markdown_rows = []
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        markdown_rows.append("| " + " | ".join(header_cells) + " |")
        markdown_rows.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            markdown_rows.append("| " + " | ".join(cells) + " |")
        
        table_text = "\n".join(markdown_rows)
        page_estimate = max(1, char_count // self.CHARS_PER_PAGE + 1)
        
        return DocumentChunk(
            content=table_text,
            chunk_type='table',
            section_title=current_section,
            section_number=current_section_number,
            page_estimate=page_estimate,
            metadata={'row_count': len(table.rows), 'col_count': len(header_cells)}
        )
    
    def _extract_figures(self, doc: Document) -> Tuple[Dict[str, bytes], Dict[str, str]]:
        figures = {}
        captions = {}
        figure_count = 0
        
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    figure_count += 1
                    figure_id = f"figure_{figure_count}"
                    figures[figure_id] = rel.target_part.blob
                except Exception as e:
                    logger.warning(f"Could not extract image {rel_id}: {e}")
        
        current_figure = 0
        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ""
            text = para.text.strip()
            
            if 'caption' in style_name and text:
                current_figure += 1
                figure_id = f"figure_{current_figure}"
                captions[figure_id] = text
            elif text.lower().startswith('figure') and ':' in text:
                current_figure += 1
                figure_id = f"figure_{current_figure}"
                captions[figure_id] = text
        
        return figures, captions
    
    def _extract_title(self, doc: Document, chunks: List[DocumentChunk]) -> Optional[str]:
        if doc.core_properties.title:
            return doc.core_properties.title
        
        for chunk in chunks:
            if chunk.chunk_type == 'heading' and chunk.heading_level in [0, 1]:
                return chunk.content
        
        return None
